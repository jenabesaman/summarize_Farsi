import os
from docx import Document
def converting(input_text):
    script_dir = os.path.dirname(os.path.abspath(__file__))
    file_path = "text.txt"
    output_path = os.path.join(script_dir, file_path)
    with open(output_path, "w") as file:
        file.write(input_text)


#
# def ConvertToTXT():
#     doc = Document("text.docx")
#
#     # Extract text and save it to a text file
#     with open("output.txt", "w", encoding="utf-8") as text_file:
#         for paragraph in doc.paragraphs:
#             text_file.write(paragraph.text + "\n")
#
# ConvertToTXT()
#
# from docx import Document
#
# # Load the Word document
# doc = Document("text.docx")
#
# # Create a text string to store the extracted text
# text_string = ""
# for paragraph in doc.paragraphs:
#     text_string += paragraph.text + " "
#
# # Remove line breaks
# text_string = text_string.replace("\n", " ")
#
# # Write the text to a text file
# with open("output.txt", "w", encoding="utf-8") as text_file:
#     text_file.write(text_string)


# import os
# import re
# import docx
# import PyPDF2
# import openpyxl
#
# def convert_file_to_text(input_file, output_file):
#     # Get the file extension
#     file_extension = os.path.splitext(input_file)[1].lower()
#
#     if file_extension == '.docx':
#         # Handle Word documents
#         doc = docx.Document(input_file)
#         with open(output_file, 'w', encoding='utf-8') as text_file:
#             for paragraph in doc.paragraphs:
#                 text_file.write(paragraph.text + ' ')
#
#     elif file_extension == '.pdf':
#         # Handle PDF documents
#         pdf_file = open(input_file, 'rb')
#         pdf_reader = PyPDF2.PdfFileReader(pdf_file)
#         with open(output_file, 'w', encoding='utf-8') as text_file:
#             for page_num in range(pdf_reader.numPages):
#                 page = pdf_reader.getPage(page_num)
#                 text = page.extractText()
#                 # Remove line breaks
#                 text = re.sub(r'\n', ' ', text)
#                 text_file.write(text)
#
#     elif file_extension in ('.xls', '.xlsx'):
#         # Handle Excel files
#         wb = openpyxl.load_workbook(input_file, data_only=True)
#         with open(output_file, 'w', encoding='utf-8') as text_file:
#             for sheet in wb.sheetnames:
#                 ws = wb[sheet]
#                 for row in ws.iter_rows(values_only=True):
#                     for cell in row:
#                         text_file.write(str(cell) + ' ')
#                     text_file.write('\n')
#
#     else:
#         raise ValueError("Unsupported file format")

# Example usage
# input_file = 'text.docx'
# output_file = 'output.txt'
# convert_file_to_text(input_file, output_file)
